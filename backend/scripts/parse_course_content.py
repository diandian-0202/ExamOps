"""
parse_course_content.py
-----------------------
Parses lecture PDFs and exam PDFs/DOCXs from backend/exam+lecture/,
chunks the text, and writes seed_chunks.sql for D1 import.

Each chunk is tagged with:
  - class_id  : which course (1 = EECS 485, 2 = EECS 370)
  - source    : 'lecture' or 'exam'
  - topic_tag : lecture topic extracted from filename (e.g. 'Networking')

Usage:
    pip install pypdf python-docx
    python3 scripts/parse_course_content.py

Output:
    backend/seed_chunks.sql  (overwritten each run)

To add a new course or file:
    1. Drop PDF/DOCX into backend/exam+lecture/
    2. Add an entry to EXAM_FILES or adjust LECTURE_PREFIX below
    3. Re-run this script, then:
         npm run db:seed          (local)
         npm run db:seed:remote   (Cloudflare)
"""

import re
import os
import sys

try:
    import pypdf
except ImportError:
    sys.exit("Missing pypdf — run: pip install pypdf")

try:
    import docx
except ImportError:
    sys.exit("Missing python-docx — run: pip install python-docx")

# ── Configuration ─────────────────────────────────────────────────────────────

COURSE_DIR = os.path.join(os.path.dirname(__file__), '..', 'exam+lecture')
OUTPUT     = os.path.join(os.path.dirname(__file__), '..', 'seed_chunks.sql')

# Classes: id → name
CLASSES = {
    1: 'EECS 485',
    2: 'EECS 370',
}

# Lecture PDFs are auto-detected by filename pattern: NN_Topic_Name.pdf
# They all belong to EECS 485 (class_id = 1)
LECTURE_CLASS_ID = 1

# Exam files: (class_id, filename)
EXAM_FILES = [
    (1, 'eecs485f25_final_solutions.pdf'),
    (1, 'eecs485f25_midterm_solutions.pdf'),
    (1, 'eecs485w25_final_solutions.pdf'),
    (1, 'eecs485w25_midterm_solutions.pdf'),
    (2, '370 F24 Final Solutions v2 (Public).docx'),
    (2, '370Midterm F24 - Answer Key.docx'),
]

# ── Extraction ────────────────────────────────────────────────────────────────

def extract_pdf(path):
    reader = pypdf.PdfReader(path)
    return '\n'.join(page.extract_text() or '' for page in reader.pages)

def extract_docx(path):
    doc = docx.Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return '\n'.join(parts)

def extract(path):
    if path.endswith('.pdf'):
        return extract_pdf(path)
    elif path.endswith('.docx'):
        return extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {path}")

# ── Cleaning ──────────────────────────────────────────────────────────────────

def clean(text):
    text = re.sub(r'uniqname:.*', '', text)
    text = re.sub(r'Licensed under.*', '', text)
    text = re.sub(r'https?://\S+', '', text)
    text = re.sub(r'Page \d+ of \d+', '', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

# ── Chunking (sliding window) ─────────────────────────────────────────────────

def sliding_chunks(text, size=500, step=350, min_len=80):
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = ' '.join(words[i:i + size])
        if len(chunk) >= min_len:
            chunks.append(chunk)
        i += step
    return chunks

# ── SQL helpers ───────────────────────────────────────────────────────────────

def esc(s):
    return s.replace("'", "''")

def sql_val(v):
    if v is None:
        return 'NULL'
    return f"'{esc(str(v))}'"

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    all_chunks = []  # list of (class_id, source, topic_tag, content)

    # ── 1. Auto-detect lecture PDFs (pattern: NN_Topic_Name.pdf) ─────────────
    lecture_pattern = re.compile(r'^(\d+)[_ ](.+)\.pdf$')
    lecture_files = []
    for fname in sorted(os.listdir(COURSE_DIR)):
        m = lecture_pattern.match(fname)
        if m:
            topic_tag = m.group(2).replace('_', ' ').replace('-', ' ').strip()
            lecture_files.append((fname, topic_tag))

    print(f"Found {len(lecture_files)} lecture PDFs:")
    for fname, tag in lecture_files:
        path = os.path.join(COURSE_DIR, fname)
        print(f"  [{tag}] {fname} ...", end=' ')
        raw = clean(extract(path))
        chunks = sliding_chunks(raw)
        for c in chunks:
            all_chunks.append((LECTURE_CLASS_ID, 'lecture', tag, c))
        print(f"{len(chunks)} chunks")

    # ── 2. Exam files ─────────────────────────────────────────────────────────
    print(f"\nProcessing {len(EXAM_FILES)} exam files:")
    for class_id, fname in EXAM_FILES:
        path = os.path.join(COURSE_DIR, fname)
        if not os.path.exists(path):
            print(f"  WARNING: not found, skipping — {fname}")
            continue
        print(f"  [exam] {fname} ...", end=' ')
        raw = clean(extract(path))
        chunks = sliding_chunks(raw)
        for c in chunks:
            all_chunks.append((class_id, 'exam', None, c))
        print(f"{len(chunks)} chunks")

    # ── 3. Write SQL ──────────────────────────────────────────────────────────
    lines = [
        '-- Auto-generated by scripts/parse_course_content.py',
        '-- DO NOT edit manually — re-run the script instead.',
        '',
    ]

    class_values = ', '.join(f"({cid}, '{name}')" for cid, name in CLASSES.items())
    lines.append(f"INSERT OR IGNORE INTO classes (id, name) VALUES {class_values};")
    lines.append('')
    lines.append('DELETE FROM course_chunks;')
    lines.append('')

    lecture_count = sum(1 for c in all_chunks if c[1] == 'lecture')
    exam_count    = sum(1 for c in all_chunks if c[1] == 'exam')

    for class_id, source, topic_tag, content in all_chunks:
        lines.append(
            f"INSERT INTO course_chunks (class_id, source, topic_tag, content) "
            f"VALUES ({class_id}, '{source}', {sql_val(topic_tag)}, '{esc(content)}');"
        )

    sql = '\n'.join(lines)
    with open(OUTPUT, 'w', encoding='utf-8') as f:
        f.write(sql)

    print(f"\nDone.")
    print(f"  Lecture chunks : {lecture_count}")
    print(f"  Exam chunks    : {exam_count}")
    print(f"  Total          : {len(all_chunks)}")
    print(f"  Written to     : {os.path.abspath(OUTPUT)}")
    print()
    print("Next steps:")
    print("  Local:      npm run db:seed")
    print("  Cloudflare: npm run db:seed:remote")

if __name__ == '__main__':
    print(f"Course dir: {os.path.abspath(COURSE_DIR)}")
    print(f"Output:     {os.path.abspath(OUTPUT)}")
    print()
    main()
